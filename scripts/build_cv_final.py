"""Build the final CV DOCX by augmenting the reformatted v1.1 CV with content
restored from the original CV (Resume_Saeed Ahmed Khan-2(1).pdf).

This script preserves the design / typography / styling of the v1.1 reformatted
CV and inserts/extends content. It does not touch the source DOCX; output is
written to Saeed_Ahmed_Khan_CV_Final.docx in the repo root.
"""

from __future__ import annotations

import copy
import os
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

REPO_ROOT = Path(__file__).resolve().parent.parent
SRC = REPO_ROOT / "Saeed_Ahmed_Khan_CV_v1_1.docx"
DST = REPO_ROOT / "Saeed_Ahmed_Khan_CV_Final.docx"


def find_paragraph(doc: Document, predicate) -> Paragraph:
    for p in doc.paragraphs:
        if predicate(p):
            return p
    raise LookupError("Paragraph not found")


def insert_paragraph_after(paragraph: Paragraph, text: str, style_template: Paragraph | None = None) -> Paragraph:
    """Insert a new paragraph after `paragraph`, copying formatting from `style_template`.

    If style_template is given we deep-copy its XML and replace the run text;
    that preserves bullet formatting, fonts, and any custom indent.
    """
    parent = paragraph._p.getparent()
    if style_template is not None:
        new_p = copy.deepcopy(style_template._p)
        # remove all existing runs from the cloned paragraph
        for r in list(new_p.findall(qn("w:r"))):
            new_p.remove(r)
        # add a single run with text matching the template's first run formatting
        first_run_template = style_template._p.find(qn("w:r"))
        if first_run_template is not None:
            new_run = copy.deepcopy(first_run_template)
            for t in list(new_run.findall(qn("w:t"))):
                new_run.remove(t)
            t_el = new_run.makeelement(qn("w:t"), {qn("xml:space"): "preserve"})
            t_el.text = text
            new_run.append(t_el)
            new_p.append(new_run)
        else:
            new_para = Paragraph(new_p, paragraph._parent)
            new_para.add_run(text)
        paragraph._p.addnext(new_p)
        return Paragraph(new_p, paragraph._parent)

    # no template: build a fresh paragraph
    new_p = paragraph._p.makeelement(qn("w:p"), {})
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = paragraph.style
    new_para.add_run(text)
    return new_para


def chain_insert_after(anchor: Paragraph, items: list[str], template: Paragraph) -> Paragraph:
    """Insert `items` as new paragraphs immediately after `anchor`, each using
    `template` for formatting. Returns the last inserted paragraph."""
    current = anchor
    for text in items:
        current = insert_paragraph_after(current, text, template)
    return current


def main() -> int:
    if not SRC.exists():
        print(f"ERROR: source CV not found at {SRC}", file=sys.stderr)
        return 1

    doc = Document(str(SRC))

    # ---- Anchor: "Drive internationalization..." (last bullet of Dean) ----
    dean_last_bullet = find_paragraph(
        doc,
        lambda p: p.text.strip().startswith("Drive internationalization"),
    )

    extra_dean_bullets = [
        "Lead curriculum development, OBE implementation, and Pakistan Engineering Council (PEC) accreditation cycles for engineering programmes.",
        "Coordinate research direction, postgraduate programme growth, faculty hiring, and laboratory and infrastructure planning.",
        "Engage statutory bodies, industry advisory boards, and external academic partners on faculty strategy and quality systems.",
    ]
    chain_insert_after(dean_last_bullet, extra_dean_bullets, dean_last_bullet)

    # ---- Anchor: "Established faculty/staff training..." (last bullet of QEC) ----
    qec_last_bullet = find_paragraph(
        doc,
        lambda p: p.text.strip().startswith("Established faculty/staff training"),
    )

    extra_qec_bullets = [
        "Coordinated programme self-assessment reports for ME / MS / PhD programmes through the QEC Programme Review Committee.",
        "Designed institutional quality benchmarks, performance indicators, and evidence portfolios aligned with the HEC QAA / QEC framework.",
        "Worked with university statutory bodies (Senate, Syndicate, Academic Council, Board of Advanced Studies and Research, Board of Faculty) as a member and convener on quality and academic matters.",
        "Liaised with HEC, Sindh HEC, and external accreditation reviewers; coordinated accreditation and re-accreditation processes.",
        "Provided leadership and direction to the QEC team, managing budget, audits, training, and reporting cycles.",
    ]
    last_qec = chain_insert_after(qec_last_bullet, extra_qec_bullets, qec_last_bullet)

    # ---- Insert a NEW role block: Secretary, SIBA Testing Services (STS) ----
    # We need to insert title-line (Normal) + date-line (Normal) + bullets (List Paragraph)
    # immediately AFTER the last QEC bullet, BEFORE the "Research Professor / Postdoctoral Fellow" block.
    # We'll pattern after how other titles look. The cleanest is to clone existing title/date paragraphs.

    rp_title = find_paragraph(
        doc,
        lambda p: p.text.strip().startswith("Research Professor / Postdoctoral Fellow"),
    )
    rp_date = find_paragraph(
        doc,
        lambda p: p.text.strip().startswith("MAY 2022") or p.text.strip().startswith("May 2022"),
    )

    # Insert blank Normal spacer (mimic spacing pattern between roles)
    blank = insert_paragraph_after(last_qec, "", style_template=last_qec)
    # Reset blank style to Normal (not List Paragraph)
    blank_p = blank._p
    # remove paragraph properties that make it a list item
    pPr = blank_p.find(qn("w:pPr"))
    if pPr is not None:
        for tag in ("w:numPr", "w:ind"):
            el = pPr.find(qn(tag))
            if el is not None:
                pPr.remove(el)

    # Title paragraph - clone from the Dean/QEC/RP title style
    title_para = insert_paragraph_after(
        blank,
        "Secretary, SIBA Testing Services (STS)   Sukkur IBA University · Pakistan",
        style_template=rp_title,
    )

    # Date paragraph - clone from the existing date style
    date_para = insert_paragraph_after(
        title_para,
        "NOVEMBER 2024 – FEBRUARY 2025",
        style_template=rp_date,
    )

    # Bullets - clone from List Paragraph style
    sts_bullets = [
        "Oversaw operations of SIBA Testing Services (STS) — the university's standardized testing arm — supporting test administration, governance, and reporting.",
        "Coordinated test scheduling, candidate administration, and result governance.",
        "Liaised with university leadership and external test-using institutions on STS operations.",
    ]
    chain_insert_after(date_para, sts_bullets, qec_last_bullet)

    # ---- Strengthen Postdoc bullets ----
    postdoc_last_bullet = find_paragraph(
        doc,
        lambda p: p.text.strip().startswith("Co-authored multiple high-impact publications in Polymers"),
    )

    extra_postdoc_bullets = [
        "Strengthened cross-border research linkages between Sukkur IBA, Korean groups (SKKU), and earlier Chinese / European collaborators.",
        "Brought back research expertise that informed Sukkur IBA's flexible electronics, soft materials, and self-powered sensing capability on return.",
    ]
    chain_insert_after(postdoc_last_bullet, extra_postdoc_bullets, postdoc_last_bullet)

    # ---- Strengthen Associate Professor bullets ----
    ap_last_bullet = find_paragraph(
        doc,
        lambda p: p.text.strip().startswith("Served as Chief Editor"),
    )

    extra_ap_bullets = [
        "Active in Boards of Studies, OBE Accreditation Committee, Programme Review Committee, and Departmental Strategic Committee.",
        "Subject Expert (Electronics) for academic and research selection committees, Sukkur IBA University.",
    ]
    chain_insert_after(ap_last_bullet, extra_ap_bullets, ap_last_bullet)

    # ---- Add additional governance / committee items (some were missing) ----
    gov_last = find_paragraph(
        doc,
        lambda p: p.text.strip().startswith("Subject Expert (Electronics)"),
    )

    extra_gov = [
        "Member, statutory bodies of Sukkur IBA University: Board of Studies, Board of Faculty, Academic Council, Board of Advanced Studies and Research (BASR), Syndicate, and Senate.",
        "Member, Departmental and Faculty Council, Faculty of Electrical Engineering, Sukkur IBA University — undergraduate and postgraduate curriculum review.",
    ]
    chain_insert_after(gov_last, extra_gov, gov_last)

    # ---- Add an Internal Notes (verification) paragraph at end of body? ----
    # Per instructions, internal-only notes belong outside the public CV.
    # We do not add anything visible here.

    doc.save(str(DST))
    print(f"Wrote {DST}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
