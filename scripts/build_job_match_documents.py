from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("_job_search_outputs")
TODAY = date(2026, 5, 9)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def apply_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in [
        ("Title", 22, "14324A"),
        ("Subtitle", 11, "5B6770"),
        ("Heading 1", 15, "14324A"),
        ("Heading 2", 13, "245B77"),
        ("Heading 3", 11, "14324A"),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(4)


def add_header_footer(section, label):
    header = section.header
    p = header.paragraphs[0]
    p.text = label
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.runs[0].font.name = "Arial"
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(91, 103, 112)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Page ")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    for r in p.runs:
        r.font.name = "Arial"
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(91, 103, 112)


def setup_doc(title):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    apply_styles(doc)
    add_header_footer(section, title)
    return doc


def add_field_table(doc, rows, widths=(2300, 7060)):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = ""
        for p in cells[0].paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.color.rgb = RGBColor(20, 50, 74)
        set_cell_shading(cells[0], "EEF4F7")
        if isinstance(value, tuple) and value[0] == "link":
            p = cells[1].paragraphs[0]
            add_hyperlink(p, value[1], value[2])
        elif isinstance(value, list):
            for i, item in enumerate(value):
                p = cells[1].paragraphs[0] if i == 0 else cells[1].add_paragraph()
                p.style = doc.styles["Normal"]
                p.text = item
        else:
            cells[1].text = str(value)
    set_table_widths(table, list(widths))
    return table


PROFILE_SUMMARY = {
    "name": "Prof. Saeed Ahmed Khan",
    "headline": "Dean of Engineering and Technology; flexible electronics, sensors, energy harvesting, thermogalvanic hydrogels, nanogenerators, and engineering education leadership.",
    "evidence": [
        "18+ years across academia, applied research, quality assurance, accreditation, and industrial electrical/instrumentation practice.",
        "60+ peer-reviewed publications and about 1,000 Google Scholar citations, including ACS Sensors, Nano Energy, Chemical Engineering Journal, Journal of Power Sources, and Journal of Materials Chemistry C.",
        "PhD on flexible sensors based on carbon nanomaterials; postdoctoral/research professor work at SKKU on EHD jet printing, PDMS devices, and triboelectric nanogenerators.",
        "Dean-level and QEC leadership, including OBE, accreditation, rankings, program review, faculty training, and institutional evidence portfolios.",
    ],
}


REWRITTEN_PROMPT = """You are Codex working locally inside D:\\dr-saeed-website. Do not deploy, change the live website, or read the whole repository. Read only the CV located in public\\documents and, if the public website cannot be opened, read only the local website content files needed to understand the profile. Build a senior-career European job-match dossier for Prof. Saeed Ahmed Khan.

Use current live sources to identify 25 to 30 open positions across Europe and Central Europe, giving priority to Tier 1 academic professor/lecturer roles that match the profile, then Tier 2 academic research roles, then Tier 3 industrial or applied R&D roles. For each position, verify that it is open as of May 9, 2026 or clearly mark the open-ended status. Record the title, institution/company, country, tier, posting date if available, deadline/open status, source link, expected salary or salary estimate, relevance percentage, brief description, application feasibility, required documents, and tailored recommendations for application templates and document content.

Create two polished DOCX files. The first file should be a detailed job-match dossier with each position starting on a separate page. The second file should list institutions that invite spontaneous, unsolicited, open, or expression-of-interest applications even when no matching advertised position is available, with recommended documents and targeting notes. Render and visually check the documents before delivery if the local renderer is available."""


jobs = [
    {
        "tier": "Tier 1 - Professor/Lecturer",
        "title": "Professor / Assistant Professor (Tenure Track) of Next-Generation Electrical and Optical Devices",
        "org": "ETH Zurich",
        "country": "Switzerland",
        "location": "Zurich",
        "posted": "2026-04-07",
        "deadline": "2026-05-31",
        "status": "Open on 2026-05-09; advertised deadline 2026-05-31.",
        "link": "https://academicpositions.com/ad/eth-zurich/2026/professor-assistant-professor-tenure-track-of-next-generation-electrical-and-optical-devices/247160",
        "salary": "Not stated. Expected Swiss faculty range: approx. CHF 130,000-230,000 gross/year depending on rank and appointment.",
        "relevance": 96,
        "description": "Faculty role in experimental materials and device physics for photonics, electronics, quantum technologies, electromechanical response, synthesis, fabrication, and characterization.",
        "fit": "Very strong match to flexible sensors, nanomaterial devices, EHD/printing, device characterization, self-powered sensing, and senior teaching/research leadership.",
        "feasibility": "High for thematic fit; competitive for ETH-level research independence. Strongest if framed as flexible, self-powered, thermogalvanic and nanocarbon device platforms with translational links.",
        "documents": ["Academic CV", "publication list", "research statement", "teaching statement", "leadership statement", "three key achievement descriptions", "highest degree certificate"],
        "recommendation": "Use a full professor/tenure-track application package: 2-page cover letter, 4-5 page research vision, 2-page teaching statement, 1-2 page leadership/diversity statement, and a selected-publications sheet showing independent contribution.",
    },
    {
        "tier": "Tier 1 - Professor/Lecturer",
        "title": "Professor in Materials Science specialising in Integrated Computational Materials Engineering",
        "org": "KTH Royal Institute of Technology",
        "country": "Sweden",
        "location": "Stockholm",
        "posted": "2026-04-22",
        "deadline": "2026-05-13",
        "status": "Open on 2026-05-09; urgent deadline 2026-05-13.",
        "link": "https://academicpositions.com/ad/kth-royal-institute-of-technology/2026/professor-in-materials-science-specialising-in-integrated-computational-materials-engineering/247702",
        "salary": "Monthly pay; not stated. Expected KTH professor range: approx. SEK 70,000-95,000/month, negotiated by merit.",
        "relevance": 78,
        "description": "Permanent professor role in materials science with ICME strategy for research and education.",
        "fit": "Moderate-high fit through materials science, device materials, fabrication/characterization, and engineering education. Computational materials focus is adjacent rather than central.",
        "feasibility": "Medium. The case needs a computational/materials strategy around soft/flexible sensing materials, thermogalvanic gels, nanocarbon composites, and data-assisted characterization.",
        "documents": ["CV", "publication list", "research and teaching portfolio", "leadership/service evidence", "certificates", "selected publications"],
        "recommendation": "Prepare a professor portfolio emphasizing program leadership, PhD/MS supervision, quality systems, and a materials roadmap from carbon nanomaterials to self-powered sensing devices.",
    },
    {
        "tier": "Tier 1 - Professor/Lecturer",
        "title": "Professor/Associate Professor (Tenure Track), Technical Physics",
        "org": "University of Eastern Finland",
        "country": "Finland",
        "location": "Kuopio/Joensuu listing",
        "posted": "2026-04-10",
        "deadline": "2026-05-11",
        "status": "Open on 2026-05-09; urgent deadline 2026-05-11.",
        "link": "https://academicpositions.com/ad/university-of-eastern-finland/2026/professorassociate-professor-tenure-track-technical-physics/247355",
        "salary": "Not stated. Expected Finnish associate/professor range: approx. EUR 5,300-8,000 gross/month based on rank and performance component.",
        "relevance": 82,
        "description": "Professor/associate professor in technical physics with links to medical physics, computational engineering, materials engineering, and industrial applications.",
        "fit": "Strong match to applied physics, materials, sensors, health monitoring, and engineering education; less specific to flexible electronics but broad enough to support the profile.",
        "feasibility": "Medium-high. Best positioned around technical physics of wearable/self-powered sensing and thermoelectric hydrogel devices.",
        "documents": ["CV", "list of publications", "teaching merits", "research plan", "leadership/service portfolio", "degree certificates"],
        "recommendation": "Use a tenure-track professor dossier with research program sections: materials-to-device physics, health/HMI sensing, energy conversion, and curriculum development in applied physics.",
    },
    {
        "tier": "Tier 1 - Professor/Lecturer",
        "title": "Assistant Professor of Ship Systems and Marine Engineering",
        "org": "Tallinn University of Technology",
        "country": "Estonia",
        "location": "Tallinn",
        "posted": "2026-04",
        "deadline": "2026-05-31",
        "status": "Open on 2026-05-09; advertised deadline 2026-05-31.",
        "link": "https://academicpositions.com/ad/tallinn-university-of-technology/2026/assistant-professor-of-ship-systems-and-marine-engineering/246084",
        "salary": "Not stated. Expected Estonia assistant professor range: approx. EUR 3,000-5,000 gross/month, depending on seniority and contract.",
        "relevance": 64,
        "description": "Assistant professor role strengthening green maritime technology, integrated ship systems, autonomous/unmanned technology, electrical, power, control, and safety systems.",
        "fit": "Partial match through electrical systems, control, instrumentation, sustainable energy, and engineering leadership; lower direct match to nanomaterials.",
        "feasibility": "Medium-low. Apply only if willing to pivot profile toward energy systems, instrumentation, control, and autonomous marine systems.",
        "documents": ["Motivation letter", "education certificates", "CV and publication list", "academic portfolio", "future vision/action plan up to 1500 words", "additional evidence"],
        "recommendation": "Use a teaching-and-systems template. Emphasize industrial instrumentation, sustainable energy project work, control/instrumentation teaching, and faculty leadership.",
    },
    {
        "tier": "Tier 1 - Professor/Lecturer",
        "title": "Assistant-Associate Professor in Biomechanics of Regenerative Implants",
        "org": "Eindhoven University of Technology",
        "country": "Netherlands",
        "location": "Eindhoven",
        "posted": "2026-05-06",
        "deadline": "2026-05-31",
        "status": "Open on 2026-05-09; advertised deadline 2026-05-31.",
        "link": "https://academicpositions.com/ad/eindhoven-university-of-technology/2026/assistant-associate-professor-in-biomechanics-of-regenerative-implants/248268",
        "salary": "Scale 11-14 stated; expected approx. EUR 4,728-7,986 gross/month depending on appointment level.",
        "relevance": 69,
        "description": "Permanent academic position at the intersection of regenerative implants, biomechanics, computational modeling, machine learning, and data science.",
        "fit": "Moderate match via wearable/health sensing, biomedical instrumentation, gel/hydrogel devices, and machine-learning-assisted sensing.",
        "feasibility": "Medium. Stronger if framed around smart implant monitoring, soft materials, self-powered biosensing, and signal/biomechanics integration.",
        "documents": ["Cover letter", "CV with publications and references", "scientific interests/plans", "teaching goals/experience", "degree evidence"],
        "recommendation": "Prepare a biomedical engineering version of the CV and a 2-page research concept on self-powered soft sensors for regenerative-device monitoring.",
    },
    {
        "tier": "Tier 1 - Professor/Lecturer",
        "title": "Associate or Full Professor in Process Systems Engineering",
        "org": "Eindhoven University of Technology",
        "country": "Netherlands",
        "location": "Eindhoven",
        "posted": "2026-03-17",
        "deadline": "2026-05-20",
        "status": "Open on 2026-05-09; advertised deadline 2026-05-20.",
        "link": "https://academicpositions.nl/ad/eindhoven-university-of-technology/2026/associate-or-full-professor-in-process-systems-engineering/246227",
        "salary": "Not stated on snippet. Expected TU/e associate/full professor range: approx. EUR 6,000-9,600 gross/month depending on scale and rank.",
        "relevance": 72,
        "description": "Permanent senior faculty position in process systems engineering, with research, teaching, supervision, and academic leadership expectations.",
        "fit": "Moderate-high through systems engineering, energy systems, instrumentation, process modeling, and leadership, but less direct for nanomaterials.",
        "feasibility": "Medium. A strong application should connect sensor-material process design, energy harvesting systems, and engineering leadership.",
        "documents": ["Cover letter", "CV with publications and three references", "scientific interests/plans", "teaching goals and experience"],
        "recommendation": "Use a senior professor template focused on systems-level contributions: sensor fabrication process chains, lab capability building, and curriculum leadership.",
    },
    {
        "tier": "Tier 1 - Professor/Lecturer",
        "title": "Research Professor - open to all scientific fields (Open BOFZAP)",
        "org": "KU Leuven",
        "country": "Belgium",
        "location": "Leuven",
        "posted": "2026-04-03",
        "deadline": "2026-09-01",
        "status": "Open on 2026-05-09; advertised deadline 2026-09-01.",
        "link": "https://www.kuleuven.be/personeel/jobsite/jobs/60640907?lang=en",
        "salary": "Not stated. Expected Belgium ZAP professor range: approx. EUR 80,000-130,000 gross/year depending on rank; advert notes BOF research-professor funding and start-up support.",
        "relevance": 88,
        "description": "Full-time junior or senior research professor posts open across fields, with 10-year BOF-funded research emphasis and limited teaching initially.",
        "fit": "Excellent strategic route because the profile can propose a full research program in flexible, self-powered, thermogalvanic and wearable sensor systems.",
        "feasibility": "High concept fit; competitive across all fields. Requires a sharp, internationally differentiated research agenda and KU Leuven faculty host alignment.",
        "documents": ["Academic CV", "research plan", "teaching statement", "publication record", "grant/leadership evidence", "host/faculty alignment material"],
        "recommendation": "Prepare a 5-year BOFZAP research program with work packages, ERC/MSCA grant pathway, PhD pipeline, and synergies with ESAT/materials/biomedical engineering.",
    },
    {
        "tier": "Tier 1 - Professor/Lecturer",
        "title": "Assistant Professor in Systems & Data",
        "org": "TU Delft",
        "country": "Netherlands",
        "location": "Delft",
        "posted": "2026-04",
        "deadline": "2026-05-10",
        "status": "Open on 2026-05-09; deadline 2026-05-10, immediate action required.",
        "link": "https://academicpositions.co.uk/ad/technische-universiteit-delft-tu-delft/2026/assistant-professor-in-systems-data/247459",
        "salary": "Advert states EUR 4,728-6,433 gross/month.",
        "relevance": 63,
        "description": "Assistant professor in data-intensive systems for science and society; permanent/academic trajectory with research and teaching.",
        "fit": "Indirect but plausible through sensor data, biomedical signal processing, HMI, and data pipelines from QEC/ranking work.",
        "feasibility": "Medium-low unless the application is positioned around data-intensive sensing infrastructure rather than materials fabrication.",
        "documents": ["CV", "cover letter", "research statement", "teaching statement", "selected publications", "degree certificate"],
        "recommendation": "Only pursue with a focused systems-data story: intelligent sensing pipelines, embedded health monitoring, and reproducible sensor data infrastructure.",
    },
    {
        "tier": "Tier 1 - Professor/Lecturer",
        "title": "Lecturer",
        "org": "Sheffield Hallam University",
        "country": "United Kingdom",
        "location": "Sheffield",
        "posted": "2026-05-05",
        "deadline": "2026-05-17",
        "status": "Open on 2026-05-09; advertised deadline 2026-05-17.",
        "link": "https://www.jobs.ac.uk/job/DRL079/lecturer",
        "salary": "Advert states GBP 39,906-44,746 per annum.",
        "relevance": 82,
        "description": "Lecturer role in engineering and built environment, including electrical/electronic engineering, mechatronics, power systems/machines, analogue/digital electronics, control and automation.",
        "fit": "Strong teaching fit with courses delivered in electronics, circuits, instrumentation, control, signals, and engineering systems.",
        "feasibility": "High for teaching breadth; salary/rank may be below current seniority. Useful as a UK teaching-entry option.",
        "documents": ["Academic CV", "cover letter", "teaching statement", "evidence of teaching and curriculum development", "degree certificates", "references"],
        "recommendation": "Use a teaching-focused CV and cover letter. Lead with electrical/electronics teaching breadth, OBE/accreditation, student supervision, and industry-grounded examples.",
    },
    {
        "tier": "Tier 1 - Professor/Lecturer",
        "title": "Lecturer in Electrical Engineering",
        "org": "Regent College London",
        "country": "United Kingdom",
        "location": "London",
        "posted": "2026-04-10",
        "deadline": "2026-06-09",
        "status": "Open on 2026-05-09; advertised deadline 2026-06-09.",
        "link": "https://www.jobs.ac.uk/job/DRD765/lecturer-in-electrical-engineering",
        "salary": "Advert states up to GBP 50,000 per annum subject to skills and experience.",
        "relevance": 78,
        "description": "Permanent lecturer role designing, developing, and delivering engineering modules in a student-focused higher education setting.",
        "fit": "Strong fit for broad electrical/electronics teaching and academic leadership; weaker research intensity than profile.",
        "feasibility": "High, especially if the goal includes UK teaching roles. Below current seniority, but straightforward match.",
        "documents": ["CV", "cover letter", "teaching philosophy", "degree certificates", "evidence of module design", "references"],
        "recommendation": "Prepare a concise teaching application: curriculum design, OBE, accreditation, student outcomes, and examples from electronics, control, sensors, and instrumentation.",
    },
    {
        "tier": "Tier 1 - Professor/Lecturer",
        "title": "Lecturer / Senior Lecturer in Electrical and Electronic Engineering",
        "org": "University of Greenwich",
        "country": "United Kingdom",
        "location": "Medway",
        "posted": "2026-05-01",
        "deadline": "2026-05-28",
        "status": "Open on 2026-05-09; advertised deadline 2026-05-28.",
        "link": "https://www.jobs.ac.uk/job/share/?job_ref=DRK307",
        "salary": "Advert states GBP 38,784-56,535 per annum.",
        "relevance": 84,
        "description": "Lecturer/senior lecturer role in electrical and electronic engineering in the School of Engineering.",
        "fit": "Strong teaching/research fit with electronics, sensors, circuits, control, instrumentation, and academic governance.",
        "feasibility": "High for senior lecturer consideration if the portfolio is tailored; salary/rank should be checked before applying.",
        "documents": ["CV", "cover letter", "teaching and research statement", "publication list", "degree certificates", "references"],
        "recommendation": "Use a UK senior lecturer template showing teaching leadership, curriculum ownership, external engagement, publications, supervision, and quality assurance.",
    },
    {
        "tier": "Tier 1 - Professor/Lecturer",
        "title": "Lecturer or Senior Lecturer in Electrical Machines",
        "org": "University of Manchester",
        "country": "United Kingdom",
        "location": "Manchester",
        "posted": "2026-04-20",
        "deadline": "2026-05-14",
        "status": "Open on 2026-05-09; urgent deadline 2026-05-14.",
        "link": "https://www.jobs.ac.uk/search/?academicDisciplineFacet%5B0%5D=engineering-and-technology&academicDisciplineFacet%5B1%5D=architecture-building-and-planning&pageSize=25&sortOrder=2&startIndex=101&subDisciplineFacet%5B0%5D=civil-engineering&subDisciplineFacet%5B1%5D=mechanical-engineering&subDisciplineFacet%5B2%5D=electrical-and-electronic-engineering&subDisciplineFacet%5B3%5D=architecture-and-building",
        "salary": "Advert states Lecturer GBP 47,389-58,225; Senior Lecturer GBP 59,966-71,566 per annum.",
        "relevance": 76,
        "description": "Teaching and research post in electrical machines within Electrical and Electronic Engineering.",
        "fit": "Good electrical engineering and energy-systems fit, though less direct to flexible electronics and nanomaterials.",
        "feasibility": "Medium-high if framed through energy harvesting, instrumentation, power/energy teaching, and control systems.",
        "documents": ["CV", "cover letter", "research statement", "teaching statement", "publication list", "degree certificates"],
        "recommendation": "Use a senior lecturer package that shows energy harvesting research, electrical engineering courses, and leadership in program quality.",
    },
    {
        "tier": "Tier 1 - Professor/Lecturer",
        "title": "Lecturer in Analogue and Digital Electronics",
        "org": "University of Manchester",
        "country": "United Kingdom",
        "location": "Manchester",
        "posted": "2026-04-20",
        "deadline": "2026-05-15",
        "status": "Open on 2026-05-09; advertised deadline 2026-05-15.",
        "link": "https://www.jobs.ac.uk/search/?academicDisciplineFacet%5B0%5D=computer-sciences&academicDisciplineFacet%5B1%5D=engineering-and-technology&distance=30&locality%5B0%5D=Manchester&location=Manchester&locationCoords%5B0%5D=53.4807593%2C-2.2426305&pageSize=25&salaryBandFacet%5B0%5D=40000-49999&salaryBandFacet%5B1%5D=50000-59999&salaryBandFacet%5B2%5D=30000-39999&salaryBandFacet%5B3%5D=60000-69999&sortOrder=2&startIndex=1&subDisciplineFacet%5B0%5D=computer-science&subDisciplineFacet%5B1%5D=software-engineering&subDisciplineFacet%5B2%5D=artificial-intelligence&subDisciplineFacet%5B3%5D=mechanical-engineering&subDisciplineFacet%5B4%5D=electrical-and-electronic-engineering",
        "salary": "Advert states GBP 47,389-58,225 per annum depending on experience.",
        "relevance": 86,
        "description": "Lecturer post in analogue and digital electronics in a major EEE department.",
        "fit": "Very strong teaching fit through electronic devices/circuits, analogue/digital electronics, VLSI, instrumentation, sensors, and circuit analysis.",
        "feasibility": "High. The profile is senior enough; the application must explain why a lecturer role is strategically desirable.",
        "documents": ["CV", "cover letter", "teaching statement", "research statement", "selected publications", "degree certificates"],
        "recommendation": "Use a teaching-research hybrid application and include a one-page module portfolio: sensors technology, EDC, VLSI, instrumentation, analogue/digital electronics.",
    },
    {
        "tier": "Tier 2 - Academic Researcher",
        "title": "Postdoc: Inverse design of broadband multi-parameter nanophotonic sensors",
        "org": "AMOLF",
        "country": "Netherlands",
        "location": "Amsterdam",
        "posted": "2026-04-11",
        "deadline": "2026-05-28",
        "status": "Open on 2026-05-09; advertised deadline 2026-05-28.",
        "link": "https://academicpositions.com/ad/amolf/2026/postdoc-inverse-design-of-broadband-multi-parameter-nanophotonic-sensors/247365",
        "salary": "Not stated. Expected Dutch postdoc range: approx. EUR 3,378-5,331 gross/month depending on scale/experience.",
        "relevance": 88,
        "description": "Postdoctoral research on inverse design of nanophotonic structures for broadband multiparameter fiber sensing with industrial collaborators.",
        "fit": "Strong sensor and nanotechnology fit; computational photonics is adjacent to the profile's materials/device fabrication strength.",
        "feasibility": "High technical match, but rank is junior relative to current dean/associate professor profile.",
        "documents": ["CV", "cover letter", "publication list", "research fit statement", "degree certificate", "references"],
        "recommendation": "Apply only if seeking a research pivot. Emphasize sensing publications, COMSOL/ANSYS capability, signal processing, and industry-linked sensor development.",
    },
    {
        "tier": "Tier 2 - Academic Researcher",
        "title": "Postdoc in Nanoparticle Waveguide Microscopy",
        "org": "Chalmers University of Technology",
        "country": "Sweden",
        "location": "Gothenburg",
        "posted": "2026-04-16",
        "deadline": "2026-05-31",
        "status": "Open on 2026-05-09; advertised deadline 2026-05-31.",
        "link": "https://academicpositions.com/ad/chalmers-university-of-technology/2026/postdoc-in-nanoparticle-waveguide-microscopy/247516",
        "salary": "Not stated. Expected Swedish postdoc range: approx. SEK 38,000-45,000/month.",
        "relevance": 80,
        "description": "Two-year postdoc, extendable, on surface-sensitive label-free waveguide scattering microscopy for nanoparticle characterization.",
        "fit": "Good match to nanomaterials, sensors, characterization, optics, image/signal processing and biomedical applications.",
        "feasibility": "Medium-high technically; lower rank than profile. Could be useful for entering a top Nordic lab.",
        "documents": ["CV", "personal letter", "degree certificate", "publication list", "references/contact details"],
        "recommendation": "Use a research-only CV version highlighting SEM/XRD/device characterization, nanomaterials, sensor publications, and health-monitoring applications.",
    },
    {
        "tier": "Tier 2 - Academic Researcher",
        "title": "Postdoctoral Research Fellow (Advanced Microelectronics Packaging)",
        "org": "Tampere University",
        "country": "Finland",
        "location": "Tampere",
        "posted": "2026-04-13",
        "deadline": "2026-05-11",
        "status": "Open on 2026-05-09; urgent deadline 2026-05-11.",
        "link": "https://academicpositions.com/ad/tampere-university/2026/postdoctoral-research-fellow-advanced-microelectronics-packaging/247383",
        "salary": "Not stated. Expected Finland postdoctoral fellow range: approx. EUR 3,700-5,000 gross/month depending on performance level.",
        "relevance": 86,
        "description": "Four-year funded postdoctoral role in heterogeneous integration for electronics-photonics co-packaging, low-TRL packaging, and subgroup leadership.",
        "fit": "Strong match to electronic packaging, printed electronics, PDMS/EHD device work, and electrical/electronics background.",
        "feasibility": "High technical fit, but title is postdoctoral. Stronger if positioned as senior postdoctoral leadership rather than entry-level postdoc.",
        "documents": ["CV", "cover letter", "publication list", "degree certificate", "references", "research fit statement"],
        "recommendation": "Highlight electronic packaging teaching, EHD printing, soft/flexible device fabrication, and ability to mentor junior researchers.",
    },
    {
        "tier": "Tier 2 - Academic Researcher",
        "title": "Post-doctoral researcher in deep learning - 24-month fixed-term contract",
        "org": "IMT Atlantique",
        "country": "France",
        "location": "Nantes/Brest project context",
        "posted": "2026-04-16",
        "deadline": "2026-05-31",
        "status": "Open on 2026-05-09; advertised deadline 2026-05-31.",
        "link": "https://academicpositions.com/ad/imt-atlantique/2026/post-doctoral-researcher-in-deep-learning-24-month-fixed-term-contract/247491",
        "salary": "Not stated. Expected French public postdoc range: approx. EUR 2,600-3,800 gross/month.",
        "relevance": 71,
        "description": "Postdoc to develop robust sensor-fusion deep learning for IoT monitoring of reusable packaging and embedded sensor models.",
        "fit": "Moderate fit through sensor fusion, IoT monitoring, Python/ML, and signal processing; less direct to materials synthesis.",
        "feasibility": "Medium. Apply if interested in AI/sensor-fusion direction rather than materials research.",
        "documents": ["CV", "cover letter", "publication list", "PhD certificate", "references"],
        "recommendation": "Prepare a compact AI-for-sensing statement referencing deep-learning-assisted respiratory monitoring and bio-signal/health monitoring work.",
    },
    {
        "tier": "Tier 2 - Academic Researcher",
        "title": "Senior Research Fellow/Research Fellow in Integrated MEMS/NEMS-Photonics Platform for Secure Communication",
        "org": "University of Southampton",
        "country": "United Kingdom",
        "location": "Southampton",
        "posted": "2026-04-21",
        "deadline": "2026-05-11",
        "status": "Open on 2026-05-09; urgent deadline 2026-05-11.",
        "link": "https://www.jobs.ac.uk/search/?academicDisciplineFacet%5B0%5D=engineering-and-technology&pageSize=25&sortOrder=1&startIndex=101&subDisciplineFacet%5B0%5D=aerospace-engineering&subDisciplineFacet%5B1%5D=electrical-and-electronic-engineering",
        "salary": "Advert states GBP 38,784-48,822 per annum.",
        "relevance": 87,
        "description": "Research fellow/senior research fellow role in integrated MEMS/NEMS-photonics platform development.",
        "fit": "Strong fit to micro/nano devices, sensors, electronic materials, photonics, fabrication, and characterization.",
        "feasibility": "High technical fit; rank may be below current seniority. Senior Research Fellow framing is better than Research Fellow.",
        "documents": ["CV", "cover letter", "publication list", "research statement", "degree certificate", "references"],
        "recommendation": "Use a device-platform CV emphasizing nanofabrication, EHD/PDMS, flexible electronics, sensors, and device characterization.",
    },
    {
        "tier": "Tier 2 - Academic Researcher",
        "title": "Research Fellow - Machine Learning for Nanoscale Semiconductor Manufacturing",
        "org": "University of Southampton",
        "country": "United Kingdom",
        "location": "Southampton",
        "posted": "2026-04-17",
        "deadline": "2026-05-14",
        "status": "Open on 2026-05-09; advertised deadline 2026-05-14.",
        "link": "https://www.jobs.ac.uk/search/?academicDisciplineFacet%5B0%5D=physical-and-environmental-sciences&academicDisciplineFacet%5B1%5D=engineering-and-technology&distance=30&locality%5B0%5D=Southampton&location=Southampton&locationCoords%5B0%5D=50.9105468%2C-1.4049018&pageSize=25&salaryBandFacet%5B0%5D=30000-39999&salaryBandFacet%5B1%5D=40000-49999&salaryBandFacet%5B2%5D=25000-29999&salaryBandFacet%5B3%5D=50000-59999&sortOrder=2&startIndex=1&subDisciplineFacet%5B0%5D=materials-science&subDisciplineFacet%5B1%5D=electrical-and-electronic-engineering",
        "salary": "Advert states GBP 36,636-44,746 per annum.",
        "relevance": 82,
        "description": "Research fellow role in machine learning for nanoscale semiconductor manufacturing within sustainable electronic technologies.",
        "fit": "Good match to semiconductor characterization, nanomaterials, EHD/printing, sensors, and Python/data methods.",
        "feasibility": "Medium-high. Needs a stronger ML/manufacturing framing than the CV currently foregrounds.",
        "documents": ["CV", "cover letter", "publication list", "degree certificate", "technical skills statement", "references"],
        "recommendation": "Add a one-page technical annex on MATLAB/Python, characterization workflows, and data-driven sensor/device process optimization.",
    },
    {
        "tier": "Tier 2 - Academic Researcher",
        "title": "Research Fellow in Nanodevices to Study Molecular Motors",
        "org": "University of Southampton",
        "country": "United Kingdom",
        "location": "Southampton",
        "posted": "2026-04-16",
        "deadline": "2026-05-13",
        "status": "Open on 2026-05-09; advertised deadline 2026-05-13.",
        "link": "https://www.jobs.ac.uk/search/?academicDisciplineFacet%5B0%5D=biological-sciences&academicDisciplineFacet%5B1%5D=physical-and-environmental-sciences&academicDisciplineFacet%5B2%5D=engineering-and-technology&distance=30&locality%5B0%5D=Southampton&location=Southampton&locationCoords%5B0%5D=50.9105468%2C-1.4049018&pageSize=25&salaryBandFacet%5B0%5D=30000-39999&salaryBandFacet%5B1%5D=25000-29999&salaryBandFacet%5B2%5D=40000-49999&sortOrder=3&startIndex=1&subDisciplineFacet%5B0%5D=biology&subDisciplineFacet%5B1%5D=other-biological-sciences&subDisciplineFacet%5B2%5D=chemistry&subDisciplineFacet%5B3%5D=physics-and-astronomy&subDisciplineFacet%5B4%5D=other-physical-sciences&subDisciplineFacet%5B5%5D=chemical-engineering&subDisciplineFacet%5B6%5D=biotechnology&subDisciplineFacet%5B7%5D=other-engineering",
        "salary": "Advert states GBP 36,636-44,746 per annum.",
        "relevance": 74,
        "description": "Research fellow role in nanodevices for studying molecular motors within digital health and biomedical engineering.",
        "fit": "Moderate-good through nanodevices, biosensing, hydrogels, smart contact lens/health-monitoring work, and device characterization.",
        "feasibility": "Medium. Biology/molecular motor angle is a stretch but nanodevice and health sensor overlap is real.",
        "documents": ["CV", "cover letter", "publication list", "research fit statement", "degree certificate", "references"],
        "recommendation": "Use a biomedical nanodevice narrative linking thermogalvanic hydrogel sensors, smart contact lens work, and self-powered physiological monitoring.",
    },
    {
        "tier": "Tier 2 - Academic Researcher",
        "title": "Research Fellow - Smart Electronic Materials & Systems",
        "org": "University of Southampton",
        "country": "United Kingdom",
        "location": "Southampton",
        "posted": "2026-04-20",
        "deadline": "2026-05-18",
        "status": "Open on 2026-05-09; advertised deadline 2026-05-18.",
        "link": "https://www.jobs.ac.uk/search/?academicDisciplineFacet%5B0%5D=biological-sciences&academicDisciplineFacet%5B1%5D=physical-and-environmental-sciences&academicDisciplineFacet%5B2%5D=engineering-and-technology&distance=30&locality%5B0%5D=Southampton&location=Southampton&locationCoords%5B0%5D=50.9105468%2C-1.4049018&pageSize=25&salaryBandFacet%5B0%5D=30000-39999&salaryBandFacet%5B1%5D=25000-29999&salaryBandFacet%5B2%5D=40000-49999&sortOrder=3&startIndex=1&subDisciplineFacet%5B0%5D=biology&subDisciplineFacet%5B1%5D=other-biological-sciences&subDisciplineFacet%5B2%5D=chemistry&subDisciplineFacet%5B3%5D=physics-and-astronomy&subDisciplineFacet%5B4%5D=other-physical-sciences&subDisciplineFacet%5B5%5D=chemical-engineering&subDisciplineFacet%5B6%5D=biotechnology&subDisciplineFacet%5B7%5D=other-engineering",
        "salary": "Advert states GBP 36,636-41,064 per annum.",
        "relevance": 89,
        "description": "Research fellow role in smart electronic materials and systems.",
        "fit": "Very strong match to flexible electronics, smart materials, sensors, nanogenerators, and printed/soft electronic systems.",
        "feasibility": "High technically. Seniority mismatch should be handled by explaining targeted UK research transition.",
        "documents": ["CV", "cover letter", "publication list", "research plan/fit note", "degree certificates", "references"],
        "recommendation": "Lead with selected ACS Sensors/Nano Energy/Chemical Engineering Journal outputs and a concise device-fabrication skills matrix.",
    },
    {
        "tier": "Tier 2 - Academic Researcher",
        "title": "Postdoc in sustainable energy and hydrogen technologies",
        "org": "Eindhoven University of Technology",
        "country": "Netherlands",
        "location": "Eindhoven",
        "posted": "2026-04-22",
        "deadline": "2026-05-22",
        "status": "Open on 2026-05-09; advertised deadline 2026-05-22.",
        "link": "https://academicpositions.it/ad/eindhoven-university-of-technology/2026/postdoc-in-sustainable-energy-and-hydrogen-technologies/247743?locale=da",
        "salary": "Scale 10 stated; expected approx. EUR 3,378-5,331 gross/month.",
        "relevance": 70,
        "description": "Postdoc in sustainable energy/hydrogen technologies with complete application package and online screening.",
        "fit": "Moderate match through energy harvesting, sustainable energy project history, and engineering systems; less direct than sensor roles.",
        "feasibility": "Medium. Best as energy-systems pivot rather than core flexible electronics position.",
        "documents": ["Cover letter", "CV with publications and three references", "list of up to five best publications"],
        "recommendation": "Use an energy-focused CV variant emphasizing hybrid nanogenerators, sustainable energy project work, and energy conversion publications.",
    },
    {
        "tier": "Tier 2 - Academic Researcher",
        "title": "Research Assistant/Fellow in Digital Twin for Advanced Manufacturing",
        "org": "Brunel University London",
        "country": "United Kingdom",
        "location": "Uxbridge, London",
        "posted": "2026-04-27",
        "deadline": "2026-05-25",
        "status": "Open on 2026-05-09; advertised deadline 2026-05-25.",
        "link": "https://www.jobs.ac.uk/job/DRI785/research-assistant-fellow-in-digital-twin-for-advanced-manufacturing-16864",
        "salary": "Advert states GBP 36,640-44,179, with potential progression to GBP 52,067.",
        "relevance": 67,
        "description": "Fixed-term research role in digital twins for advanced manufacturing at BCAST.",
        "fit": "Moderate fit through manufacturing, modeling, sensors, data, and engineering systems; less direct to flexible electronics.",
        "feasibility": "Medium. Apply if interested in process optimization and digital manufacturing rather than academic faculty track.",
        "documents": ["CV", "cover letter", "degree certificate", "publication/technical output list", "references"],
        "recommendation": "Frame the application around sensor/process monitoring, manufacturing digitalization, and experience with MATLAB/Python/COMSOL/ANSYS.",
    },
    {
        "tier": "Tier 2 - Academic Researcher",
        "title": "Research Associate in Aircraft-Based Power Electronics and Power Systems Design (Hardware lead)",
        "org": "University of Strathclyde",
        "country": "United Kingdom",
        "location": "Glasgow",
        "posted": "2026-05-01",
        "deadline": "2026-05-17",
        "status": "Open on 2026-05-09; advertised deadline 2026-05-17.",
        "link": "https://www.jobs.ac.uk/job/DKY235",
        "salary": "Advert snippet states GBP 37,694-46,049 per annum.",
        "relevance": 71,
        "description": "Research associate role in aircraft-based power electronics and power systems design.",
        "fit": "Good electrical engineering, hardware, power and systems fit; lower match to nanomaterials/flexible sensors.",
        "feasibility": "Medium. Stronger if positioned around power/energy systems, electronic devices, and applied engineering experience.",
        "documents": ["CV", "cover letter", "publication list", "technical skills statement", "degree certificate", "references"],
        "recommendation": "Use a power-electronics hardware template emphasizing electronics, instrumentation, energy harvesting, and industrial electrical engineering roles.",
    },
    {
        "tier": "Tier 3 - Industrial/Applied R&D",
        "title": "R&D Nano Engineer",
        "org": "SRON - Space Research Organization Netherlands",
        "country": "Netherlands",
        "location": "Leiden",
        "posted": "2026-05-07 approx.",
        "deadline": "2026-05-10",
        "status": "Open on 2026-05-09; urgent deadline 2026-05-10.",
        "link": "https://academicpositions.com/jobs/field/micro-nanotechnology",
        "salary": "Not stated in snippet. Expected Dutch research engineer range: approx. EUR 3,500-5,800 gross/month.",
        "relevance": 82,
        "description": "Cleanroom nano-fabrication role for space detectors including deposition, etch, lithography, metrology and process documentation.",
        "fit": "Strong hands-on match to thin films, CVD/EB-PVD, SEM/XRD, nanofabrication and device characterization.",
        "feasibility": "Medium-high technically, but may prefer candidates with intensive cleanroom operations and local work authorization.",
        "documents": ["Industry CV", "cover letter", "skills matrix", "certificates", "references"],
        "recommendation": "Use an industry R&D CV: one-page technical skills matrix, fabrication/characterization tools, project deliverables, and quality/process documentation experience.",
    },
    {
        "tier": "Tier 3 - Industrial/Applied R&D",
        "title": "Senior Research Engineer - Nitride Thin-Film Materials (Photonics Application)",
        "org": "Silicon Austria Labs",
        "country": "Austria",
        "location": "Graz",
        "posted": "2026-04 approx.",
        "deadline": "Open/unspecified in source snippet",
        "status": "Appeared active in current AcademicPositions results on 2026-05-09; confirm before applying.",
        "link": "https://academicpositions.com/jobs/field/micro-nanotechnology",
        "salary": "Not stated in snippet. Expected Austrian senior R&D engineer range: approx. EUR 55,000-75,000 gross/year, depending on collective agreement and experience.",
        "relevance": 84,
        "description": "Develop AlN thin films for photonic integrated circuits using sputtering/CVD, optical and structural characterization, process optimization in cleanroom, publications and IP.",
        "fit": "Strong match to thin films, CVD/EB-PVD, SEM, XRD, electronic materials, and photonic/electronic device research.",
        "feasibility": "High technical fit if industrial cleanroom and thin-film process development are emphasized.",
        "documents": ["Industry CV", "cover letter", "tool/process matrix", "publication/patent list", "degree certificates", "references"],
        "recommendation": "Prepare a SAL-specific industry-research package with thin-film, characterization, device fabrication, and translational publication evidence.",
    },
    {
        "tier": "Tier 3 - Industrial/Applied R&D",
        "title": "Microfluidics Modeling Specialist",
        "org": "imec",
        "country": "Belgium",
        "location": "Leuven",
        "posted": "Current in 2026 AcademicPositions result",
        "deadline": "Open/unspecified in source snippet",
        "status": "Appeared active in current AcademicPositions results on 2026-05-09; confirm in imec portal before applying.",
        "link": "https://academicpositions.com/jobs/field/micro-nanotechnology",
        "salary": "Not stated. Expected imec specialist range: approx. EUR 55,000-85,000 gross/year depending on level.",
        "relevance": 74,
        "description": "CFD/multiphysics simulation for microfluidic liquid cooling and life-science systems using ANSYS Fluent, parameterized workflows and R&D collaboration.",
        "fit": "Moderate-high match through ANSYS/COMSOL, sensor systems, micro/nano devices, health applications and engineering modeling.",
        "feasibility": "Medium-high if simulation depth is demonstrable; may require stronger microfluidics portfolio.",
        "documents": ["Industry CV", "cover letter", "simulation portfolio", "degree certificates", "references"],
        "recommendation": "Create a technical portfolio page showing COMSOL/ANSYS use, sensor/device modeling, and experimental validation examples.",
    },
    {
        "tier": "Tier 3 - Industrial/Applied R&D",
        "title": "Wet Process Development Coordinator",
        "org": "imec",
        "country": "Belgium",
        "location": "Leuven",
        "posted": "2026-05 approx.",
        "deadline": "Open/unspecified in source snippet",
        "status": "Appeared active in current AcademicPositions results on 2026-05-09; confirm in imec portal before applying.",
        "link": "https://academicpositions.com/jobs/field/micro-nanotechnology",
        "salary": "Not stated. Expected senior semiconductor process role: approx. EUR 60,000-90,000 gross/year depending on experience.",
        "relevance": 76,
        "description": "Coordinate wet process quality initiatives and module maturity for logic/3D/memory projects; DOE/SPC/FDC, tech transfer and customer support.",
        "fit": "Good process/quality match via QEC leadership, evidence portfolios, fabrication and characterization; less direct wet processing evidence.",
        "feasibility": "Medium. Strong if positioned around quality systems plus micro/nano fabrication experience.",
        "documents": ["Industry CV", "cover letter", "quality/process improvement evidence", "tool matrix", "certificates"],
        "recommendation": "Use a hybrid technical-quality CV that combines fabrication tools with accreditation/quality-system leadership and process documentation.",
    },
    {
        "tier": "Tier 3 - Industrial/Applied R&D",
        "title": "Project Engineer - Electrical Engineering and Automation",
        "org": "Institut Laue-Langevin",
        "country": "France",
        "location": "Grenoble",
        "posted": "2026-05 approx.",
        "deadline": "Open/unspecified in snippet",
        "status": "Appeared active in current AcademicPositions results on 2026-05-09; confirm before applying.",
        "link": "https://academicpositions.com/jobs/field/electrical-engineering-engineering",
        "salary": "Not stated. Expected France/Grenoble engineering role: approx. EUR 45,000-65,000 gross/year plus benefits.",
        "relevance": 72,
        "description": "Permanent project engineering role managing electrical/automation projects in a regulated nuclear research facility from requirements to commissioning.",
        "fit": "Good fit to industrial electrical/instrumentation background, control systems, project coordination and quality processes.",
        "feasibility": "Medium-high if project engineering and regulated-systems experience are emphasized.",
        "documents": ["Industry CV", "cover letter", "project list", "certificates", "references"],
        "recommendation": "Prepare a project-based CV with electrical maintenance/instrumentation roles, control systems teaching, and QEC process/audit leadership.",
    },
    {
        "tier": "Tier 3 - Industrial/Applied R&D",
        "title": "Research Engineer in AI accelerator compiler integration and design space exploration - startup preparation",
        "org": "KU Leuven",
        "country": "Belgium",
        "location": "Leuven",
        "posted": "2026-05 approx.",
        "deadline": "2026-05-10",
        "status": "Open on 2026-05-09; urgent deadline 2026-05-10.",
        "link": "https://academicpositions.com/jobs/field/electrical-engineering-engineering",
        "salary": "Not stated. Expected KU Leuven research engineer range: approx. EUR 3,500-5,500 gross/month.",
        "relevance": 61,
        "description": "Research engineer role linking design-space exploration tools to AI accelerator compiler backends, with possible spin-off continuation.",
        "fit": "Indirect match through electronics, programming, Python/C++, systems and HCI; weak for materials.",
        "feasibility": "Medium-low unless hardware/software co-design experience is prominent.",
        "documents": ["Technical CV", "cover letter", "software/hardware project portfolio", "degree certificates", "references"],
        "recommendation": "Only apply if tailoring toward embedded systems. Highlight Python/C++, electronic systems, and research-to-startup collaboration experience.",
    },
    {
        "tier": "Tier 3 - Industrial/Applied R&D",
        "title": "Senior Researcher Digital Health Computing f/m/d",
        "org": "St. Polten University of Applied Sciences",
        "country": "Austria",
        "location": "St. Polten",
        "posted": "2026-05-01",
        "deadline": "Unspecified",
        "status": "Open/unspecified deadline in active 2026 listing; confirm before applying.",
        "link": "https://academicpositions.com/ad/st-polten-university-applied-sciences/2026/senior-researcher-digital-health-computing-f-m-d/248131",
        "salary": "Advert states at least EUR 65,900 gross/year for 40h/week all-in.",
        "relevance": 75,
        "description": "Permanent senior researcher role in digital health computing, biomedical engineering, medical technology, AI, computer vision and HCI.",
        "fit": "Good applied-health fit through wearable health monitoring, HMI, smart contact lens/respiration/facial perception publications and signal processing.",
        "feasibility": "Medium-high if application foregrounds digital health applications and AI/signal processing rather than materials synthesis.",
        "documents": ["Research CV", "cover letter", "publication list", "project/grant portfolio", "degree certificates", "references"],
        "recommendation": "Use a digital-health CV variant with a section on self-powered physiological sensing, HMI, deep-learning-assisted monitoring, and grant leadership.",
    },
    {
        "tier": "Tier 3 - Industrial/Applied R&D",
        "title": "Operations Manager (UK National Ion Beam Centre)",
        "org": "University of Surrey / ATI Electronics",
        "country": "United Kingdom",
        "location": "Guildford",
        "posted": "2026-05-07",
        "deadline": "2026-05-27",
        "status": "Open on 2026-05-09; advertised deadline 2026-05-27.",
        "link": "https://www.jobs.ac.uk/search/?academicDisciplineFacet%5B0%5D=engineering-and-technology&academicDisciplineFacet%5B1%5D=business-and-management-studies&distance=30&locality%5B0%5D=London&locality%5B1%5D=Docklands+campus&location=Docklands+campus&locationCoords%5B0%5D=51.5072178%2C-0.1275862&locationCoords%5B1%5D=51.507369%2C0.0641766&pageSize=25&salaryBandFacet%5B0%5D=40000-49999&salaryBandFacet%5B1%5D=50000-59999&salaryBandFacet%5B2%5D=60000-69999&salaryBandFacet%5B3%5D=30000-39999&salaryBandFacet%5B4%5D=70000-79999&sortOrder=2&startIndex=51&subDisciplineFacet%5B0%5D=civil-engineering&subDisciplineFacet%5B1%5D=mechanical-engineering&subDisciplineFacet%5B2%5D=aerospace-engineering&subDisciplineFacet%5B3%5D=electrical-and-electronic-engineering&subDisciplineFacet%5B4%5D=production-engineering-and-manufacturing&subDisciplineFacet%5B5%5D=chemical-engineering&subDisciplineFacet%5B6%5D=other-engineering&subDisciplineFacet%5B7%5D=management",
        "salary": "Advert states GBP 47,389-56,535 per annum.",
        "relevance": 68,
        "description": "Operations manager role for a national ion beam facility with engineering, stakeholder and operational responsibilities.",
        "fit": "Moderate fit through research infrastructure, materials characterization, lab management, and quality systems.",
        "feasibility": "Medium. More management/infrastructure than research; good if leadership and facility operations are a goal.",
        "documents": ["Management CV", "cover letter", "operations/project portfolio", "certificates", "references"],
        "recommendation": "Use a leadership CV emphasizing dean/QEC roles, lab capability building, evidence pipelines, accreditation, and research facility governance.",
    },
]


spontaneous_targets = [
    {
        "name": "Hochschule Bielefeld - University of Applied Sciences and Arts (HSBI)",
        "country": "Germany",
        "link": "https://www.hsbi.de/en/careers/advice-on-your-application",
        "basis": "HSBI states that unsolicited applications may be sent depending on area, and separately invites contact for lecturer interests and general interest in becoming a university-of-applied-sciences professor.",
        "relevance": "High for lecturer/professor networking in electrical/electronics, sensors, mechatronics and applied engineering education.",
        "documents": "Motivation letter, CV, certificates; for professorships prepare detailed appointment package, publication list, teaching portfolio and industry/practice evidence.",
        "recommendation": "Contact the appointments officer and HR lecturer contact first. Do not send a generic email; attach a one-page fit note for electrical/electronics, sensors and OBE/accreditation teaching.",
    },
    {
        "name": "Karl Landsteiner University of Health Sciences",
        "country": "Austria",
        "link": "https://www.kl.ac.at/en/university/career-kl/application-procedure-employees",
        "basis": "The application procedure page explicitly mentions unsolicited applications and lists upload documents.",
        "relevance": "Medium-high for biomedical sensing, smart health monitoring, hydrogels and digital health, but less direct for electronics faculty.",
        "documents": "Current CV, motivation letter, proof of highest education, publication list for academic positions, references and certificates if available.",
        "recommendation": "Use a health-technology profile: self-powered physiological sensing, smart contact lens, respiratory monitoring and facial perception publications.",
    },
    {
        "name": "KU Leuven - Evolutionary Stress Ecology and Ecotoxicology Laboratory",
        "country": "Belgium",
        "link": "https://bio.kuleuven.be/eeb/rs/job-opportunities",
        "basis": "The lab states it is happy to discuss PhD and postdoctoral projects that fit broad research themes.",
        "relevance": "Low-medium for this profile because it is biological/ecotoxicology rather than EEE; useful mainly as an example of KU Leuven group-level spontaneous routes.",
        "documents": "Two-page CV, one-page motivation letter, transcripts where relevant, contact information for two referees.",
        "recommendation": "Do not prioritize unless a specific sensor/ecotoxicology collaboration is identified. Better use the KU Leuven BOFZAP route for this profile.",
    },
    {
        "name": "CCG/ZGDV Institute",
        "country": "Portugal",
        "link": "https://www.ccg.pt/en/recruitment",
        "basis": "The recruitment page has a spontaneous application route and invites candidates to submit applications.",
        "relevance": "Medium for HCI, ubiquitous computing, information engineering and applied technology transfer; not a classic professor track.",
        "documents": "CV and application form materials; for research roles also prepare cover letter, certificates and publication/portfolio evidence.",
        "recommendation": "Target HMI, sensing interfaces, digital health and applied engineering systems. Use an applied R&D CV rather than a full professor package.",
    },
    {
        "name": "IDIAP Research Institute",
        "country": "Switzerland",
        "link": "https://careers.werecruit.io/en/idiap/offers/spontaneous-application-7f85e1",
        "basis": "IDIAP provides a spontaneous application form requiring resume and cover letter and asks which group the candidate is most interested in.",
        "relevance": "Medium for human-centered health AI, activity understanding, signal processing and HMI; less direct for materials fabrication.",
        "documents": "Resume, cover letter, availability, preferred group and possible referee at Idiap.",
        "recommendation": "Use a digital-health/HMI pitch: self-powered health monitoring, respiratory and posture recognition, smart pen/contact lens applications, and signal-processing supervision.",
    },
    {
        "name": "NAVER LABS Europe",
        "country": "France",
        "link": "https://careers.werecruit.io/en/naver-labs-europe/offres/candidature-spontanee-c0650f",
        "basis": "The careers form accepts spontaneous applications and allows resume, recommendation letter, application letter and attachments.",
        "relevance": "Medium for robotics, AI, HCI, sensor systems and autonomous technologies; industrial research rather than university faculty.",
        "documents": "Resume, recommendation letter, application letter and attachments.",
        "recommendation": "Only target if preparing an industry-research pivot. Focus on HMI, smart sensing, signal processing and engineering systems, not academic administration.",
    },
    {
        "name": "ISC Paris Group",
        "country": "France",
        "link": "https://www.iscparis.com/en/recruitment-isc-paris/",
        "basis": "The recruitment page accepts unsolicited applications for temporary adjunct professor roles and lists required academic documents.",
        "relevance": "Low for engineering, but useful if targeting adjunct teaching in technology management, innovation, research methods or engineering leadership.",
        "documents": "CV, cover letter, area of expertise; for faculty roles, CV, cover letter, research statement, teaching statement and recent evaluations if available.",
        "recommendation": "Use only for adjunct/management teaching, not as a primary technical match. Frame engineering leadership, quality assurance and technology transfer.",
    },
    {
        "name": "Museum national d'Histoire naturelle doctoral school ED227",
        "country": "France",
        "link": "https://formation.mnhn.fr/en/phd/admission-2895",
        "basis": "The doctoral admission page describes a spontaneous application procedure after approval from a future HDR doctoral supervisor.",
        "relevance": "Low for Prof. Khan personally; relevant only if arranging doctoral collaboration or student placement.",
        "documents": "Supervisor approval, online application through ADUM, doctoral admission documents.",
        "recommendation": "Do not prioritize for employment. Consider only as a collaboration/student pathway if a sensor/environmental monitoring supervisor connection exists.",
    },
]


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def build_positions_doc(path):
    selected_jobs = [
        job
        for job in jobs
        if job["title"]
        not in {
            "Research Engineer in AI accelerator compiler integration and design space exploration - startup preparation",
            "Operations Manager (UK National Ion Beam Centre)",
        }
    ]

    doc = setup_doc("European Position Match Dossier")
    doc.add_heading("European Position Match Dossier", 0)
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Prepared locally for Prof. Saeed Ahmed Khan | Open-source check date: 9 May 2026")
    doc.add_paragraph(
        "Scope: 30 currently open or active Europe/Central Europe opportunities prioritized by fit to the CV and website profile. "
        "Open status, salary and document requirements should be rechecked immediately before submission, especially for urgent May deadlines."
    )

    doc.add_heading("Rewritten Codex Prompt", level=1)
    doc.add_paragraph(REWRITTEN_PROMPT)

    doc.add_heading("Profile Basis Used For Matching", level=1)
    add_field_table(
        doc,
        [
            ("Candidate", PROFILE_SUMMARY["name"]),
            ("Positioning", PROFILE_SUMMARY["headline"]),
            ("Primary evidence", PROFILE_SUMMARY["evidence"]),
        ],
    )

    doc.add_heading("Tier Summary", level=1)
    tier_counts = {}
    for job in selected_jobs:
        tier_counts[job["tier"]] = tier_counts.get(job["tier"], 0) + 1
    add_field_table(
        doc,
        [
            ("Tier 1 roles", f"{tier_counts.get('Tier 1 - Professor/Lecturer', 0)} professor/lecturer roles prioritized first."),
            ("Tier 2 roles", f"{tier_counts.get('Tier 2 - Academic Researcher', 0)} academic researcher/postdoc/senior research roles."),
            ("Tier 3 roles", f"{tier_counts.get('Tier 3 - Industrial/Applied R&D', 0)} industrial or applied R&D roles."),
            ("Best immediate targets", "ETH Zurich, KU Leuven BOFZAP, University of Greenwich, University of Manchester Analogue/Digital Electronics, University of Southampton Smart Electronic Materials, Tampere Advanced Microelectronics Packaging, Silicon Austria Labs."),
        ],
    )

    doc.add_page_break()
    for idx, job in enumerate(selected_jobs, start=1):
        doc.add_heading(f"{idx}. {job['title']}", level=1)
        doc.add_paragraph(f"{job['org']} | {job['location']}, {job['country']} | {job['tier']}")
        rows = [
            ("Relevance", f"{job['relevance']}%"),
            ("Open status", job["status"]),
            ("Posted", job["posted"]),
            ("Deadline", job["deadline"]),
            ("Expected salary", job["salary"]),
            ("Position link", ("link", "Open advert/source", job["link"])),
            ("Short description", job["description"]),
            ("Why it matches the profile", job["fit"]),
            ("Application feasibility", job["feasibility"]),
            ("Required documents", job["documents"]),
            ("Recommended template/content", job["recommendation"]),
        ]
        add_field_table(doc, rows)
        if idx != len(selected_jobs):
            doc.add_page_break()

    doc.save(path)


def build_spontaneous_doc(path):
    doc = setup_doc("Spontaneous Application Targets")
    doc.add_heading("Spontaneous / Unsolicited Application Targets", 0)
    doc.add_paragraph(
        "These are not advertised job openings. They are institutions or research-linked organisations whose pages indicate a spontaneous, unsolicited, open, or expression-of-interest application route. "
        "Use them for networking, host alignment, adjunct/visiting roles, and future calls."
    )
    add_field_table(
        doc,
        [
            ("Recommended package", "One-page target email, two-page academic CV or industry CV as appropriate, full CV, selected-publications list, research concept, teaching/leadership note, degree certificates, and references."),
            ("Best-fit targets", "HSBI, Karl Landsteiner University, IDIAP, NAVER LABS Europe, CCG/ZGDV Institute."),
            ("Caution", "Do not send a generic professor package to every target. Each spontaneous application should mention a named group, lab, department, or contact route."),
        ],
    )
    doc.add_page_break()

    for idx, target in enumerate(spontaneous_targets, start=1):
        doc.add_heading(f"{idx}. {target['name']}", level=1)
        rows = [
            ("Country", target["country"]),
            ("Application route", ("link", "Open source page", target["link"])),
            ("Basis for listing", target["basis"]),
            ("Profile relevance", target["relevance"]),
            ("Documents to prepare", target["documents"]),
            ("Recommendation", target["recommendation"]),
        ]
        add_field_table(doc, rows)
        if idx != len(spontaneous_targets):
            doc.add_page_break()

    doc.save(path)


def main():
    OUT_DIR.mkdir(exist_ok=True)
    build_positions_doc(OUT_DIR / "Prof_Saeed_Europe_Position_Match_Dossier.docx")
    build_spontaneous_doc(OUT_DIR / "Prof_Saeed_Spontaneous_Application_Targets.docx")


if __name__ == "__main__":
    main()
